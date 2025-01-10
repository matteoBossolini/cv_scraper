from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import json
from datetime import datetime

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ['top', 'left', 'bottom', 'right']:
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = parse_xml(r'<w:tblBorders {} w:val="{}" w:sz="{}" w:space="0" w:color="{}"/>'.format(nsdecls('w'), edge_data.get('val', 'single'), edge_data.get('sz', '4'), edge_data.get('color', '000000')))
            tcPr.append(element)

def add_formatted_paragraph(cell, text, bold=False, italic=False, font_size=11, space_after=0):
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.italic = italic
    paragraph.space_after = Pt(space_after)
    return paragraph

def format_languages(languages):
    if not languages:
        return "Not specified"
    
    formatted_languages = []
    for lang_entry in languages:
        if lang_entry.get("language") and lang_entry.get("level"):
            lang = lang_entry["language"][0]
            level = lang_entry["level"][0].capitalize()
            formatted_languages.append(f"{lang} ({level})")
    
    return ", ".join(formatted_languages) if formatted_languages else "Not specified"

def create_resume(data):
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Create main table
    main_table = doc.add_table(rows=1, cols=1)
    main_table.autofit = False
    main_table.allow_autofit = False
    
    # Remove all borders from main table
    for cell in main_table._cells:
        set_cell_border(
            cell,
            top={"sz": "0", "val": "none", "color": "auto"},
            bottom={"sz": "0", "val": "none", "color": "auto"},
            left={"sz": "0", "val": "none", "color": "auto"},
            right={"sz": "0", "val": "none", "color": "auto"},
        )

    title_cell = main_table.rows[0].cells[0]
    
    # Add title
    add_formatted_paragraph(title_cell, "Resümee", bold=True, font_size=16, space_after=6)

    # Personal information table
    info_table = title_cell.add_table(rows=5, cols=2)
    info_table.autofit = False
    info_table.allow_autofit = False
    
    # Set column widths
    for row in info_table.rows:
        row.cells[0].width = Inches(2)
        row.cells[1].width = Inches(4)
    
    info_pairs = [
        ("Name:", data["name"]),
        ("Qualifikation:", data["main_qualification"]),
        ("Wohnort:", data["current_location"]),
        ("Geburtsdatum:", data["date_of_birth"]),
        ("Staatsangehörigkeit:", data["nationality"])
    ]

    for i, (key, value) in enumerate(info_pairs):
        cells = info_table.rows[i].cells
        add_formatted_paragraph(cells[0], key, space_after=0)
        add_formatted_paragraph(cells[1], value, space_after=0)

    # Combine and sort experiences
    experiences = []
    for exp in data["work_experience"]:
        experiences.append({"type": "work", "index": exp["index"], "data": exp})
    for edu in data["education"]:
        experiences.append({"type": "education", "index": edu["index"], "data": edu})
    experiences.sort(key=lambda x: x["index"])

    current_section = None
    
    # Add minimal spacing after personal info
    title_cell.add_paragraph().space_after = Pt(6)

    for exp in experiences:
        if current_section != exp["type"]:
            current_section = exp["type"]
            # Add section header directly in title cell
            header_text = "Berufserfahrung:" if exp["type"] == "work" else "Weiterbildung:"
            header_p = title_cell.add_paragraph(header_text)
            header_p.style = 'Heading 1'
            header_run = header_p.runs[0]
            header_run.font.size = Pt(14)
            header_run.font.bold = True
            header_p.space_after = Pt(6)

        # Experience
        if exp["type"] == "work":
            # Date
            date_p = title_cell.add_paragraph()
            date_p.space_after = Pt(0)
            date_range = f"{exp['data']['date_of_start']} – {exp['data']['date_of_end']}"
            date_run = date_p.add_run(date_range)
            date_run.italic = True
            date_run.font.size = Pt(11)
            
            # Role and company
            role_p = title_cell.add_paragraph()
            role_p.space_after = Pt(2)
            role_company = f"{exp['data']['role']}, {exp['data']['company_name']}"
            if exp['data'].get('location') and exp['data']['location'] != "Not specified":
                role_company += f", {exp['data']['location']}"
            role_run = role_p.add_run(role_company)
            role_run.bold = True
            role_run.font.size = Pt(11)
            
            # Tasks
            if exp["data"].get("tasks"):
                for task in exp["data"]["tasks"]:
                    task_p = title_cell.add_paragraph()
                    task_p.space_after = Pt(2)
                    task_run = task_p.add_run(f"• {task}")
                    task_run.font.size = Pt(11)

        else:  # education
            # Date
            date_p = title_cell.add_paragraph()
            date_p.space_after = Pt(0)
            date_range = f"{exp['data']['date_of_start']} – {exp['data']['date_of_end']}"
            date_run = date_p.add_run(date_range)
            date_run.italic = True
            date_run.font.size = Pt(11)
            
            # Institution and qualification
            inst_p = title_cell.add_paragraph()
            inst_p.space_after = Pt(2)
            inst_run = inst_p.add_run(exp['data']['institution'])
            inst_run.bold = True
            inst_run.font.size = Pt(11)
            
            qual_p = title_cell.add_paragraph()
            qual_p.space_after = Pt(2)
            qual_run = qual_p.add_run(f"Abschluss: {exp['data']['diploma']}")
            qual_run.font.size = Pt(11)
            
            # Certifications
            if exp["data"].get("certifications"):
                for cert in exp["data"]["certifications"]:
                    cert_p = title_cell.add_paragraph()
                    cert_p.space_after = Pt(2)
                    cert_run = cert_p.add_run(f"• {cert}")
                    cert_run.font.size = Pt(11)

        # Add minimal spacing between entries
        title_cell.add_paragraph().space_after = Pt(4)

    # Additional information section
    if data.get("it_skills") or data.get("languages") or data.get("other_qualification"):
        # Create table for additional information
        additional_table = title_cell.add_table(rows=3, cols=2)
        for row in additional_table.rows:
            row.cells[0].width = Inches(2)
            row.cells[1].width = Inches(4)
        
        # IT Skills
        add_formatted_paragraph(additional_table.rows[0].cells[0], "EDV-Kenntnisse:", space_after=0)
        add_formatted_paragraph(additional_table.rows[0].cells[1], 
                              ", ".join(data.get("it_skills", ["Not specified"])), space_after=0)
        
        # Languages
        add_formatted_paragraph(additional_table.rows[1].cells[0], "Sprachkenntnisse:", space_after=0)
        add_formatted_paragraph(additional_table.rows[1].cells[1], 
                              format_languages(data.get("languages", [])), space_after=0)
        
        # Other qualifications
        add_formatted_paragraph(additional_table.rows[2].cells[0], "Sonstige Qualifikationen:", space_after=0)
        add_formatted_paragraph(additional_table.rows[2].cells[1], 
                              "; ".join(data.get("other_qualification", ["Not specified"])), space_after=0)

    # Save the document
    doc.save('resume.docx')


with open('data.json', 'r') as f:
    data = json.load(f)
create_resume(data)